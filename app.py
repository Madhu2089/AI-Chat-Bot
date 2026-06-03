from dotenv import load_dotenv
load_dotenv()  # Load environment variables from .env
import streamlit as st
import google.generativeai as genai
import os
import docx
import PyPDF2
from docx import Document
from docx2pdf import convert
from PIL import Image, ImageEnhance, ImageFilter
import io
import base64
import uuid
import json
import datetime
import sqlite3
from utils import compress_pdf, compress_docx, enhance_image, summarize_text, format_chat_message
from database import Database


# Page configuration
st.set_page_config(
    page_title="AI Assistant for Document Processing",
    page_icon="📄",
    layout="wide",
)

# Set your Gemini API key here
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

# Initialize database
@st.cache_resource
def get_database():
    return Database()

db = get_database()

# Initialize session state variables
if "initialized" not in st.session_state:
    # Get or create user in database
    user_id = db.get_or_create_user()
    st.session_state.user_id = user_id
    
    # Load conversations from database
    conversations = db.get_all_conversations(user_id)
    st.session_state.conversations = {}
    
    if conversations:
        # Load conversations from database
        for conv in conversations:
            conversation_data = db.get_conversation(conv['id'])
            if conversation_data:
                st.session_state.conversations[conv['id']] = {
                    "title": conversation_data['title'],
                    "created_at": conversation_data['created_at'],
                    "messages": conversation_data['messages']
                }
    
    # If no conversations exist, create a new one
    if not st.session_state.conversations:
        new_id = str(uuid.uuid4())
        st.session_state.conversations[new_id] = {
            "messages": [],
            "created_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
            "title": "New Conversation"
        }
        db.save_conversation(new_id, user_id, "New Conversation", [])
        st.session_state.current_conversation_id = new_id
    else:
        # Use the most recent conversation as current
        st.session_state.current_conversation_id = list(st.session_state.conversations.keys())[0]
    
    # Set chat history to current conversation
    st.session_state.chat_history = st.session_state.conversations[st.session_state.current_conversation_id]["messages"]
    st.session_state.initialized = True

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "gemini_api_key" not in st.session_state:
    st.session_state.gemini_api_key = GEMINI_API_KEY

# Sidebar for API configuration and tools
with st.sidebar:
    st.title("🛠️ Configuration & Tools")
    
    # API Key configuration
    api_key = st.text_input("Enter Gemini API Key:", type="password", value=st.session_state.gemini_api_key)
    if api_key:
        st.session_state.gemini_api_key = api_key
        
    st.markdown("---")
    
    # Tools section
    st.subheader("Document Tools")
    
    # Document upload with clear instructions
    st.markdown("##### Upload documents (PDF, Word, or text files)")
    st.markdown("Click below to select a file from your computer:")
    uploaded_doc = st.file_uploader("Upload Document", type=["docx", "pdf", "txt"], 
                                  help="Select PDF, Word, or text files to process")
    
    if uploaded_doc is not None:
        st.success(f"Successfully uploaded: {uploaded_doc.name}")
        file_details = {"Filename": uploaded_doc.name, "Size": f"{uploaded_doc.size / 1024:.2f} KB"}
        st.write(file_details)
    else:
        st.info("Please select a document file to upload. Click on 'Browse files' to select from your computer.")
        
        # Document processing options
        doc_option = st.selectbox(
            "Choose an action:",
            ["Export to Word", "Export to PDF", "Compress document", "Summarize text"]
        )
        
        if st.button("Process Document"):
            with st.spinner("Processing document..."):
                # Process document based on selected option
                try:
                    if doc_option == "Export to Word" and uploaded_doc.type == "application/pdf":
                        # Convert PDF to Word functionality would go here
                        # This would require more complex processing, simplified for demo
                        st.success("PDF to Word conversion would happen here")
                        
                    elif doc_option == "Export to PDF" and uploaded_doc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        # Convert Word to PDF
                        buffer = io.BytesIO()
                        doc = Document(uploaded_doc)
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            label="Download PDF",
                            data=buffer,
                            file_name=f"{uploaded_doc.name.split('.')[0]}.pdf",
                            mime="application/pdf"
                        )
                        
                    elif doc_option == "Compress document":
                        if uploaded_doc.type == "application/pdf":
                            compressed_pdf = compress_pdf(uploaded_doc)
                            st.download_button(
                                label="Download Compressed PDF",
                                data=compressed_pdf,
                                file_name=f"compressed_{uploaded_doc.name}",
                                mime="application/pdf"
                            )
                        elif uploaded_doc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            compressed_docx = compress_docx(uploaded_doc)
                            st.download_button(
                                label="Download Compressed Word Doc",
                                data=compressed_docx,
                                file_name=f"compressed_{uploaded_doc.name}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    
                    elif doc_option == "Summarize text":
                        if uploaded_doc.type == "application/pdf":
                            reader = PyPDF2.PdfReader(uploaded_doc)
                            text = ""
                            for page_num in range(len(reader.pages)):
                                text += reader.pages[page_num].extract_text()
                        elif uploaded_doc.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            doc = Document(uploaded_doc)
                            text = "\n".join([para.text for para in doc.paragraphs])
                        else:  # text file
                            text = uploaded_doc.getvalue().decode()
                        
                        summary = summarize_text(text, st.session_state.gemini_api_key)
                        st.text_area("Summary", summary, height=200)
                        
                except Exception as e:
                    st.error(f"Error processing document: {str(e)}")
    
    st.markdown("---")
    
    # Image processing with clear instructions
    st.subheader("Image Tools")
    st.markdown("##### Upload images (PNG, JPG, or JPEG)")
    st.markdown("Click below to select an image from your computer:")
    uploaded_image = st.file_uploader("Upload Image", type=["png", "jpg", "jpeg"],
                                    help="Select PNG, JPG, or JPEG files to enhance")
    
    if uploaded_image is not None:
        try:
            st.success(f"Successfully uploaded: {uploaded_image.name}")
            image = Image.open(uploaded_image)
            st.image(image, caption="Uploaded Image", use_column_width=True)
        except Exception as e:
            st.error(f"Error opening image: {str(e)}")
            st.info("Please try uploading a different image file.")
    else:
        st.info("Please select an image file to upload. Click on 'Browse files' to select from your computer.")
    
    # Only show enhancement options if an image is uploaded and successfully opened
    if uploaded_image is not None and 'image' in locals():
        enhancement_option = st.slider("Enhancement Level", 1.0, 3.0, 1.5, 0.1)
        sharpen_option = st.slider("Sharpening Level", 1.0, 3.0, 1.5, 0.1)
        
        if st.button("Enhance Image"):
            with st.spinner("Enhancing image..."):
                try:
                    enhanced_image = enhance_image(image, enhancement_option, sharpen_option)
                    st.image(enhanced_image, caption="Enhanced Image", use_column_width=True)
                    
                    # Prepare image for download
                    buf = io.BytesIO()
                    enhanced_image.save(buf, format="PNG")
                    byte_im = buf.getvalue()
                    
                    st.download_button(
                        label="Download Enhanced Image",
                        data=byte_im,
                        file_name=f"enhanced_{uploaded_image.name}",
                        mime=f"image/{uploaded_image.name.split('.')[-1]}"
                    )
                except Exception as e:
                    st.error(f"Error enhancing image: {str(e)}")
                    st.info("Please try with a different image or different enhancement settings.")

# Function to create a new chat
def create_new_chat():
    new_id = str(uuid.uuid4())
    st.session_state.current_conversation_id = new_id
    
    # Add to memory
    st.session_state.conversations[new_id] = {
        "messages": [],
        "created_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "title": "New Conversation"
    }
    
    # Save to database
    db.save_conversation(new_id, st.session_state.user_id, "New Conversation", [])
    
    # Update chat history
    st.session_state.chat_history = []
    st.rerun()

# Function to load a chat
def load_chat(conversation_id):
    st.session_state.current_conversation_id = conversation_id
    st.session_state.chat_history = st.session_state.conversations[conversation_id]["messages"]
    st.rerun()

# Function to update chat title
def update_chat_title(conversation_id, new_title):
    # Update in memory
    st.session_state.conversations[conversation_id]["title"] = new_title
    
    # Update in database
    db.save_conversation(
        conversation_id, 
        st.session_state.user_id, 
        new_title, 
        st.session_state.conversations[conversation_id]["messages"]
    )

# Function to delete a chat
def delete_chat(conversation_id):
    if conversation_id in st.session_state.conversations:
        # Delete from memory
        del st.session_state.conversations[conversation_id]
        
        # Delete from database
        db.delete_conversation(conversation_id)
        
        # If deleted current conversation, switch to another one
        if conversation_id == st.session_state.current_conversation_id:
            if st.session_state.conversations:
                # Switch to first available conversation
                new_current_id = list(st.session_state.conversations.keys())[0]
                st.session_state.current_conversation_id = new_current_id
                st.session_state.chat_history = st.session_state.conversations[new_current_id]["messages"]
            else:
                # Create a new conversation if none left
                create_new_chat()
        st.rerun()

# Main content area with two columns
col1, col2 = st.columns([1, 3])

# Left column for chat history
with col1:
    st.subheader("Chat History")
    
    # New chat button
    if st.button("➕ New Chat", key="new_chat_btn"):
        create_new_chat()
    
    # Display chat history
    st.markdown("---")
    for conv_id, conv_data in sorted(
        st.session_state.conversations.items(), 
        key=lambda x: x[1]["created_at"], 
        reverse=True
    ):
        chat_title = conv_data["title"]
        chat_date = conv_data["created_at"]
        
        # Check if this is the current conversation
        is_current = conv_id == st.session_state.current_conversation_id
        
        # Create a box for each conversation
        chat_box_style = "background-color: #e6f3ff; padding: 10px; border-radius: 5px; margin: 5px 0;" if is_current else "padding: 10px; border-radius: 5px; margin: 5px 0;"
        
        with st.container():
            st.markdown(f"<div style='{chat_box_style}'>", unsafe_allow_html=True)
            
            # Create a row with chat title button and delete button
            col_title, col_delete = st.columns([5, 1])
            
            with col_title:
                if st.button(f"📝 {chat_title}", key=f"chat_{conv_id}"):
                    load_chat(conv_id)
                st.caption(f"Created: {chat_date}")
            
            with col_delete:
                if st.button("🗑️", key=f"delete_{conv_id}", help="Delete this conversation"):
                    delete_chat(conv_id)
            
            st.markdown("</div>", unsafe_allow_html=True)
    
# Right column for chat interface
with col2:
    st.title("AI Assistant with Document Processing")
    st.markdown("""
    This application provides AI-powered assistance for document processing, including:
    - Document compression and conversion
    - Image enhancement
    - Text summarization
    - AI-powered chat assistant using Gemini API
    """)

    # AI Chat interface
    st.header("💬 AI Chat Assistant")
    
    # Current chat title
    current_chat = st.session_state.conversations[st.session_state.current_conversation_id]
    current_title = current_chat["title"]
    
    # Allow user to edit the title
    new_title = st.text_input("Chat Title:", value=current_title, key="chat_title_input")
    if new_title != current_title:
        update_chat_title(st.session_state.current_conversation_id, new_title)

    if not st.session_state.gemini_api_key:
        st.warning("Please enter your Gemini API key in the sidebar to use the chat functionality.")
    else:
        # Display current chat messages
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Get user input
        user_input = st.chat_input("Ask me anything...")
        
        if user_input:
            # Add user message to chat history
            st.session_state.chat_history.append({"role": "user", "content": user_input})
            
            # Also update the conversation storage
            st.session_state.conversations[st.session_state.current_conversation_id]["messages"] = st.session_state.chat_history
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(user_input)
            
            # Generate AI response
            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                full_response = ""
                
                try:
                    # Configure Gemini API
                    genai.configure(api_key=st.session_state.gemini_api_key)
                    
                    # Create a generative model
                    # Use the newest model name - Gemini 1.5
                    model = genai.GenerativeModel('gemini-1.5-pro')
                    
                    # Process the conversation and display the response
                    chat = model.start_chat(history=[
                        {"role": m["role"], "parts": [m["content"]]} 
                        for m in st.session_state.chat_history[:-1]  # exclude last user message
                    ])
                    
                    response = chat.send_message(user_input)
                    full_response = response.text
                    
                    message_placeholder.markdown(full_response)
                except Exception as e:
                    error_message = f"Error generating response: {str(e)}"
                    message_placeholder.error(error_message)
                    full_response = error_message
            
            # Add assistant response to chat history
            st.session_state.chat_history.append({"role": "assistant", "content": full_response})
            
            # Update conversation storage in memory
            st.session_state.conversations[st.session_state.current_conversation_id]["messages"] = st.session_state.chat_history
            
            # Save conversation to database
            current_title = st.session_state.conversations[st.session_state.current_conversation_id]["title"]
            db.save_conversation(
                st.session_state.current_conversation_id,
                st.session_state.user_id,
                current_title,
                st.session_state.chat_history
            )
            
            # If this is the first message, update the title based on content
            if len(st.session_state.chat_history) == 2:  # First Q&A pair
                # Update title based on user query
                if len(user_input) > 30:
                    auto_title = user_input[:27] + "..."
                else:
                    auto_title = user_input
                update_chat_title(st.session_state.current_conversation_id, auto_title)
                st.rerun()